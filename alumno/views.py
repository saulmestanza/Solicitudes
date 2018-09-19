# -*- coding: utf-8 -*-
from django import forms
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.views.generic import View
from django.views.generic.edit import FormView, CreateView, UpdateView
from django.views.generic.list import ListView
from django.views.generic import DeleteView, TemplateView
from django.core.files import File
from django.core.mail import EmailMessage
from django.core.urlresolvers import reverse_lazy
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin, PermissionRequiredMixin
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.views import update_session_auth_hash
from django.utils.translation import ugettext as _
from django.http import HttpResponse, HttpResponseNotFound, Http404,  HttpResponseRedirect
from django.shortcuts import render, get_object_or_404, get_list_or_404, reverse
from django.db.models import Q
from django.template import loader
from django.conf import settings
from django.forms.utils import ErrorList
from django.utils.timezone import datetime #important if using timezones
from forms import *
from models import *
from serializers import *
from profesor.models import *
from administrador.models import *

"""
from mimetypes import MimeTypes
from pyPdf import PdfFileWriter, PdfFileReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib.units import cm
"""
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import json
import random
import StringIO
import string
import os
import locale
from os import path as os_path, SEEK_SET as os_SEEK_SET

# Create your views here.
class AlumnoListView(UserPassesTestMixin, ListView):
    model = Alumno
    template_name = 'alumno/alumno_list.html'

    
    def test_func(self):
        return self.request.user.is_superuser

    
    def get_queryset(self):
        return super(AlumnoListView, self).get_queryset().all()

    
    def get_context_data(self, **kwargs):
        context = super(AlumnoListView, self).get_context_data(**kwargs)
        context['tittle'] = "Alumnos"
        context['motto'] = "Tabla para visualizar Alumnos"
        context['table_id'] = "alumno"
        return context


class AlumnoCreateView(PermissionRequiredMixin, CreateView):
    form_class = AlumnoForm
    template_name = 'alumno/alumno_create.html'
    permission_required = (
        'alumno.list_alumnos', 
        'alumno.add_alumno', 
        'alumno.change_alumno', 
        'alumno.delete_alumno',
    )

    def send_mail_to_user(self, user, request):
        context = {
            'email': user.email,
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'user': user,
            'protocol': 'https',
        }
        subject_template_name = 'alumno/new_account_subject.txt'
        email_template_name = 'alumno/new_account.html'
        subject = loader.render_to_string(subject_template_name, context)
        subject = ''.join(subject.splitlines())
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=[user.email],
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)
    
    def get_success_url(self):
        return '/alumno/alumno-list/'

    def get(self, request, next=None):
        return super(AlumnoCreateView, self).get(request)

    def form_valid(self, form):
        try:
            if User.objects.filter(username=form.instance.user.username).exists():
                errors = form._errors.setdefault("email", ErrorList())
                errors.append(u"Ya existe un usuario registrado con este email")
                response = super(AlumnoCreateView, self).form_invalid(form)
            else:
                response = super(AlumnoCreateView, self).form_valid(form)
                _user_id_ = form.instance.user.id
                user = User.objects.filter(pk=_user_id_).first()
                alumno = Alumno.objects.filter(user=user).first()
                alumno.is_active = True
                alumno.save()
                self.send_mail_to_user(user, self.request)
        except Exception, e:
            print 'Error', e
            errors = form._errors.setdefault("email", ErrorList())
            errors.append(u"Ya existe un usuario registrado con este email")
            if form.instance.user.id:
                form.instance.user.delete()
            response = super(AlumnoCreateView, self).form_invalid(form)
        return response

    def form_invalid(self, form):
        return super(AlumnoCreateView, self).form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super(AlumnoCreateView, self).get_context_data(**kwargs)
        context['tittle'] = "Alumnos"
        context['motto'] = "Formulario para crear un nuevo Alumno"
        return context


class AlumnoUpdateView(PermissionRequiredMixin, UpdateView):
    form_class = AlumnoForm
    template_name = 'alumno/alumno_edit.html'
    success_url = '/alumno/alumno-list/'
    permission_required = (
        'alumno.list_alumnos', 
        'alumno.add_alumno', 
        'alumno.change_alumno', 
        'alumno.delete_alumno',
    )

    def get_object(self, queryset=None):
        if not Alumno.objects.filter(pk=self.kwargs['id']).exists():
            raise Http404
        alumno = Alumno.objects.get(pk=self.kwargs['id'])
        if alumno.deleted:
            raise Http404
        return alumno

    def form_valid(self, form):
        clean = form.cleaned_data
        try:
            response = super(AlumnoUpdateView, self).form_valid(form)
            _user_id_ = form.instance.user.id
            user = User.objects.filter(pk=_user_id_).first()
            alumno = Alumno.objects.filter(user=user).first()
            alumno.is_active = True
            alumno.save()
        except Exception, e:
            print e
            errors = form._errors.setdefault("email", ErrorList())
            errors.append(u"Ya existe un usuario registrado con este email")
            response = super(AlumnoUpdateView, self).form_invalid(form)
        return response

    def form_invalid(self, form):
        return super(AlumnoUpdateView, self).form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super(AlumnoUpdateView, self).get_context_data(**kwargs)
        context['tittle'] = "Alumnos"
        context['motto'] = "Formulario para editar Alumnos"
        return context


class AlumnoPasswordUpdateView(PermissionRequiredMixin, FormView):
    form_class = AlumnoSetPasswordForm
    template_name = 'alumno/alumno_create.html'
    success_url = '/alumno/alumno-list/'
    permission_required = (
        'alumno.list_alumnos', 
        'alumno.add_alumno', 
        'alumno.change_alumno', 
        'alumno.delete_alumno',
    )

    def get_form(self):
        if not Alumno.objects.filter(id=self.kwargs['id']).exists():
            raise Http404
        self.alumno = Alumno.objects.get(id=self.kwargs['id'])
        user = self.alumno.user
        if self.request.method == 'GET':
            form = self.form_class(user=user)
        elif self.request.method == 'POST':
            form = self.form_class(user=user, data=self.request.POST)
        return form

    def get_success_url(self):
        if Alumno.objects.filter(id=self.kwargs['id']).exists():
            return '/alumno/alumno-edit/%s/'%self.kwargs['id']
        else:
            return self.success_url

    def form_valid(self, form):
        form = self.get_form()
        if form.is_valid():
            form.save()
            update_session_auth_hash(self.request, form.user)
        else:
            return super(AlumnoPasswordUpdateView, self).form_invalid(form)
        return super(AlumnoPasswordUpdateView, self).form_valid(form)

    def get_context_data(self, **kwargs):
        context = super(AlumnoPasswordUpdateView, self).get_context_data(**kwargs)
        context['tittle'] = "Alumnos"
        context['motto'] = u"Formulario para configurar la contraseña del Alumno %s %s"%(self.alumno.user.first_name, self.alumno.user.last_name)
        return context


class AlumnoDeleteView(PermissionRequiredMixin, DeleteView):
    model = Alumno
    success_url = '/alumno/alumno-list/'
    permission_required = (
        'alumno.list_alumnos', 
        'alumno.add_alumno', 
        'alumno.change_alumno', 
        'alumno.delete_alumno',
    )

    def get_object(self):
        id = self.kwargs['id']
        alumno = Alumno.objects.get(id=id)
        _user_ = User.objects.filter(pk=alumno.user.id)
        _user_.delete()
        return alumno

    def get(self, request, *args, **kwargs):
        return self.post(request, args, kwargs)


## SEGUIMIENTO 
class AlumnoSeguimientoListView(UserPassesTestMixin, ListView):
    model = ProcesoAlumno
    template_name = 'alumno/alumno_seguimiento.html'
    profesor = None

    def test_func(self):
        return self.request.user

    def get_queryset(self):
        if self.request.user.is_superuser or self.request.user.is_staff:
            return super(AlumnoSeguimientoListView, self).get_queryset().all()
        else:
            if Alumno.objects.filter(user__username=self.request.user.username).exists():
                return super(AlumnoSeguimientoListView, self).get_queryset().filter(alumn_id=self.request.user.alumno.id)
            else:
                if Profesor.objects.filter(user__username=self.request.user.username).exists():
                    self.profesor = Profesor.objects.filter(user__username=self.request.user.username).first()
                    return super(AlumnoSeguimientoListView, self).get_queryset().all()
                else:
                    raise Http404
    
    def get_context_data(self, **kwargs):
        context = super(AlumnoSeguimientoListView, self).get_context_data(**kwargs)
        context['tittle'] = "Seguimiento de Trámites"
        context['motto'] = "Bandeja de seguimiento"
        context['procesos'] = Proceso.objects.filter(deleted=False)
        context['table_id'] = "proceso_alumno"
        context['profesor'] = self.profesor
        context['historial'] = Historial.objects.filter(deleted=False)
        context['carreras'] = Carrera.objects.filter(deleted=False)
        context['periodo'] = Periodo.objects.filter(deleted=False).first()
        return context


class AlumnoSeguimientoItemUpdateView(PermissionRequiredMixin, UpdateView):
    model = ProcesoAlumno
    form_class = ProcesoAlumnoCompleteForm

    template_name = 'alumno/alumno_seguimiento_edit.html'
    success_url = '/alumno/alumno-seguimiento/'
    permission_required = (
        'alumno.add_procesoalumno',
    )

    def send_mail_to_user(self, alumno, profesor, proceso_alumno, historial, request):
        _user_ = User.objects.filter(is_superuser=True).first()
        _to_ = [alumno.user.email, profesor.user.email, _user_.email]
        context = {
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'protocol': 'https',
            'alumno': alumno,
            'profesor': profesor,
            'historial': historial,
            'proceso_alumno': proceso_alumno,
            'today':datetime.today(),
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'], reverse('alumno-seguimiento')),
        }
        email_template_name = 'alumno/edit_process.html'
        subject = proceso_alumno
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=_to_,
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def send_mail_to_profesor(self, alumno, profesor, proceso_alumno, nota, request):
        if settings.PRODUCTION:
            locale.setlocale(locale.LC_ALL, 'spanish')
        else:
            locale.setlocale(locale.LC_ALL, 'es_ES')

        today = datetime.today()

        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.5)

        p =document.add_picture(settings.BASE_DIR+'/header.png', width=Inches(3))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('')

        p = document.add_paragraph('Guayaquil {0:%d} de {0:%B} del {0:%Y}.'.format(today))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph('')
        document.add_paragraph('')

        document.add_paragraph('De mis consideraciones')
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        document.add_paragraph('')
        document.add_paragraph('')

        _text_ = ''
        if 'Gracia' in proceso_alumno.process.name:
            _text_ = 'el examen de gracia en la materia %s al estudiante %s'%(proceso_alumno.subject.name, alumno.user.first_name, alumno.user.last_name)
        else:
            _text_ = 'la recalificación del estudiante %s %s en los temas solicitados'%(alumno.user.first_name, alumno.user.last_name)

        body = "Yo ,%s %s , docente de la carrera %s he procedido a realizar %s."%(
            profesor.first_name, profesor.last_name, proceso_alumno.subject.cicles.carrer.name, _text_
            )
        document.add_paragraph(body)
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
        # add table ------------------
        table = document.add_table(1, 4)
        table.style = 'TableGrid'
        # populate header row --------
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Carrera'
        heading_cells[1].text = 'Materia'
        heading_cells[2].text = 'Parcial'
        heading_cells[3].text = 'Nota'

        # add a data row for each item
        cells = table.add_row().cells
        cells[0].text = proceso_alumno.subject.cicles.carrer.name
        cells[1].text = proceso_alumno.subject.name
        if 'Gracia' in proceso_alumno.process.name:
            cells[2].text = 'Examen de Gracia'
        else:
            cells[2].text = proceso_alumno.parcial
        cells[3].text = nota.nota

        document.add_paragraph('')

        document.add_paragraph('Atentamente,')

        document.add_paragraph('')

        document.add_paragraph('%s %s'%(profesor.first_name, profesor.last_name))
        document.add_paragraph('_________________________________')

        document.save(settings.BASE_DIR+'/formato_notas.docx')
        _file_ = File(open(settings.BASE_DIR+'/formato_notas.docx', 'rb'))
        _to_ = [profesor.email, alumno.user.email]
        context = {
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'protocol': 'https',
            'proceso_alumno': proceso_alumno,
            'today':datetime.today(),
        }
        email_template_name = 'alumno/register_nota_profesor.html'
        subject = proceso_alumno
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=_to_,
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.attach("formato_notas.docx", _file_.read(), 'application/docx')
        email_message.send(fail_silently=False)

    def send_mail_to_recalificador(self, alumno, profesor, proceso_alumno, request):
        _to_ = [profesor.email, alumno.user.email]
        procesoalumnoitem = ProcesoAlumnoItems.objects.filter(process_alumno_id = proceso_alumno.id)
        context = {
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'protocol': 'https',
            'alumno': alumno,
            'profesor': profesor,
            'proceso_alumno': proceso_alumno,
            'proceso_alumno_items': procesoalumnoitem,
            'today':datetime.today(),
        }
        email_template_name = 'alumno/recalificador_sign_in.html'
        subject = proceso_alumno
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=_to_,
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def test_func(self):
        return self.request.user

    def post(self, request, *args, **kwargs):
        form_class = self.get_form_class()
        form = self.get_form(form_class)
        if form.is_valid():
            return self.form_valid(form)
        else:
            print form.errors

    def form_valid(self, form):
        clean = form.cleaned_data
        description =  clean['description']
        document = clean['documento']
        try:
            response = HttpResponseRedirect('/alumno/alumno-seguimiento/')
            proceso_alumno = ProcesoAlumno.objects.filter(pk=self.kwargs['id']).first()
            proceso_alumno.status = clean['status']
            if clean['status'] == 'RE':
                proceso_alumno.status = 'IR'
            proceso_alumno.save()
            alumno = Alumno.objects.filter(pk=proceso_alumno.alumn_id).first()
            materia = Materia.objects.filter(name=proceso_alumno.subject).first()
            profesor = Profesor.objects.filter(subjects=materia).first()

            # Extras profesores
            extras = clean['extras']
            if extras:
                for profe in extras: 
                    if not (profe in proceso_alumno.extras.all()):
                        self.send_mail_to_recalificador(alumno, profe.user, proceso_alumno, self.request)
                        proceso_alumno.extras.add(profe)
                        proceso_alumno.save()

            # Crear historial
            historial = Historial.objects.create(
                process_alumno=proceso_alumno, 
                status = proceso_alumno.status_verbose(),
                description=description,
                document=document,
                created_by = ("%s %s")%(self.request.user.first_name, self.request.user.last_name)
            )
            historial.save()

            # Guardar notas
            if clean['input_nota_0']:
                nota = Nota.objects.create(
                    profesor = "%s %s"%(self.request.user.first_name, self.request.user.last_name),
                    nota = clean['input_nota_0']
                )
                nota.save()
                if not (nota in proceso_alumno.notes.all()):
                    proceso_alumno.notes.add(nota)
                    proceso_alumno.save()
                    self.send_mail_to_profesor(alumno, self.request.user, proceso_alumno, nota, self.request)

            # Enviar correo
            self.send_mail_to_user(alumno, profesor, proceso_alumno, historial, self.request)
        except Exception, e:
            print e
            errors = form._errors.setdefault("error", ErrorList())
            response = super(AlumnoSeguimientoItemUpdateView, self).form_invalid(form)
        return response

    def form_invalid(self, form):
        return super(AlumnoSeguimientoItemUpdateView, self).form_invalid(form)

    def get_object(self, queryset=None):
        if not ProcesoAlumno.objects.filter(pk=self.kwargs['id']).exists():
            raise Http404
        proceso_alumno = ProcesoAlumno.objects.get(pk=self.kwargs['id'])
        if proceso_alumno.deleted:
            raise Http404
        return proceso_alumno

    def get_context_data(self, **kwargs):
        context = super(AlumnoSeguimientoItemUpdateView, self).get_context_data(**kwargs)
        context['tittle'] = "Seguimiento de Trámites"
        context['proceso_alumno'] = self.get_object()
        context['process_alumno_items'] = ProcesoAlumnoItems.objects.filter(process_alumno=self.get_object())
        context['profesores'] = Profesor.objects.filter(deleted=False, user__is_active=True)
        context['profesor'] = Profesor.objects.filter(subjects=self.get_object().subject.id).first()
        context['motto'] = "Bandeja de seguimiento - Editar"
        return context


class ProcesoAlumnoCreateView(CreateView):
    model = ProcesoAlumno
    form_class = ProcesoAlumnoForm

    def post(self, request):
        if request.method == "POST":
            form = ProcesoAlumnoForm(request.POST)
            message = ''
            if(form.is_valid()):
                form.save()
                procesos = ProcesoAlumno.objects.latest('id')
                serializer = ProcesoAlumnoSerializer(procesos)
                return HttpResponse(json.dumps({'items': serializer.data}))
            else:
                print form.errors
                return HttpResponse(json.dumps({'message': message}))
        return HttpResponse(json.dumps({'message': message}))


class ProcesoAlumnoItemsCreateView(CreateView):
    model = ProcesoAlumnoItems
    form_class = ProcesoAlumnoItemsForm

    def post(self, request):
        if request.method == "POST":
            form = ProcesoAlumnoItemsForm(request.POST)
            message = 'something wrong!'
            # process_id = request
            if(form.is_valid()):
                form.save()
                proceos_items = ProcesoAlumnoItems.objects.latest('id')
                serializer = ProcesoAlumnoItemsSerializer(proceos_items)
                return HttpResponse(json.dumps({'items': serializer.data}))
            else:
                print form.errors
                return HttpResponse(json.dumps({'message': message}))
        return HttpResponse(json.dumps({'message': message}))


class ProcesoAlumnoItemsDeleteView(CreateView):
    model = ProcesoAlumnoItems
    form_class = ProcesoAlumnoItemsForm

    def post(self, request, id):
        if request.method == "POST":
            procesoalumnoitem = ProcesoAlumnoItems.objects.filter(pk=id)
            procesoalumnoitem.delete()
        return HttpResponse(json.dumps({'message': 'ok'}))


class AlumnoTramiteCreate(UserPassesTestMixin, CreateView):
    model = ProcesoAlumno
    fields = '__all__'
    template_name = 'alumno/alumno_tramite_create.html'
    success_url = '/alumno/alumno-seguimiento/'

    
    def test_func(self):
        if self.request.user.is_superuser:
            raise Http404
        else:
            return self.request.user.is_active

    def send_mail(self, alumno, profesor, proceso_alumno, request):
        _user_ = User.objects.filter(is_superuser=True).first()
        _to_ = [alumno.user.email, profesor.user.email, _user_.email]
        procesoalumnoitem = ProcesoAlumnoItems.objects.filter(process_alumno_id = proceso_alumno.id)
        context = {
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'protocol': 'https',
            'alumno': alumno,
            'profesor': profesor,
            'proceso_alumno': proceso_alumno,
            'proceso_alumno_items': procesoalumnoitem,
            'today':datetime.today(),
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'], reverse('alumno-seguimiento')),
        }
        email_template_name = 'alumno/create_process.html'
        subject = proceso_alumno
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=_to_,
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def post(self, request):
        if request.method == "POST":
            alumn_id = request.POST.get("alumno_id", "")
            if alumn_id:
                alumno = Alumno.objects.filter(pk=alumn_id).first()
            else:
                alumno = Alumno.objects.filter(user_id=self.request.user.id).first()
            proceso_alumno = ProcesoAlumno.objects.filter(alumn__id=alumno.id).latest('id')
            materia = Materia.objects.filter(name=proceso_alumno.subject).first()
            profesor = Profesor.objects.filter(subjects=materia).first()

            historial = Historial.objects.create(
                process_alumno = proceso_alumno,
                status = proceso_alumno.status_verbose(),
                description = u"Trámite %s Ingresado por el Alumno %s %s"%(
                    proceso_alumno.process.name,
                    alumno.user.first_name,
                    alumno.user.last_name
                ),
                created_by = "%s %s"%(alumno.user.first_name, alumno.user.last_name)
            )
            historial.save()

            self.send_mail(alumno, profesor, proceso_alumno, request)
        return HttpResponseRedirect(reverse('alumno-seguimiento'))
    
    def get_context_data(self, **kwargs):
        context = super(AlumnoTramiteCreate, self).get_context_data(**kwargs)
        context['tittle'] = "Crear Solicitud"
        context['motto'] = "Llene todos los datos"

        context['periodos'] = Periodo.objects.filter(deleted=False)
        context['procesos'] = Proceso.objects.all()
        context['ciclos'] = Ciclo.objects.all()
        context['carreras'] = Carrera.objects.all()
        context['materias'] = Materia.objects.all()

        alumno = Alumno.objects.filter(user=self.request.user).first()
        context['alumno'] = alumno

        form = context['form']
        for field in form.fields:
            form.fields[field].widget.attrs['class'] = 'form-control'
            if form.fields[field].help_text:
                form.fields[field].widget.attrs['data-toggle'] = 'tooltip'
                form.fields[field].widget.attrs['data-placement'] = 'top'
                form.fields[field].widget.attrs['title'] = form.fields[field].help_text
        return context
   

## Historial List y Edit 


class HistorialCreateView(UserPassesTestMixin, CreateView):
    form_class = HistorialForm
    success_url = '/alumno/alumno-seguimiento/'
    template_name = 'alumno/alumno_historial.html'

    def send_mail(self, alumno, profesor, proceso_alumno, historial, request):
        _user_ = User.objects.filter(is_superuser=True).first()
        _to_ = [alumno.user.email, profesor.user.email, _user_.email]
        context = {
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'protocol': 'https',
            'alumno': alumno,
            'user': self.request.user,
            'profesor': profesor,
            'historial': historial,
            'proceso_alumno': proceso_alumno,
            'today':datetime.today(),
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'], reverse('alumno-seguimiento')),
        }
        email_template_name = 'alumno/edit_historial.html'
        subject = proceso_alumno
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=_to_,
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def test_func(self):
        return self.request.user.is_active

    def form_valid(self, form):
        clean = form.cleaned_data
        is_ok =  clean['is_ok']
        try:
            context = self.get_context_data()
            self.object = form.save()
            response = super(HistorialCreateView, self).form_valid(form)
            
            historial = Historial.objects.filter(pk=form.instance.id).first()
            proceso_alumno = ProcesoAlumno.objects.filter(pk=self.kwargs['id']).first()
            proceso_alumno.is_ok = is_ok
            # Recalificacion
            if not 'Gracia' in proceso_alumno.process.name:
                if is_ok and proceso_alumno.status == 'ID':
                    proceso_alumno.status = 'AE'
                elif not is_ok and proceso_alumno.status == 'ID':
                    proceso_alumno.status = 'RE'
            proceso_alumno.save()
            historial.status = proceso_alumno.status_verbose()
            historial.save()
            alumno = proceso_alumno.alumn
            materia = Materia.objects.filter(name=proceso_alumno.subject).first()
            profesor = Profesor.objects.filter(subjects=materia).first()
            self.send_mail(alumno, profesor, proceso_alumno, historial, self.request)

        except Exception as e:
            print e
            errors = form._errors.setdefault("name", ErrorList())
            response = super(HistorialCreateView, self).form_invalid(form)
        return response

    def form_invalid(self, form):
        return super(HistorialCreateView, self).form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super(HistorialCreateView, self).get_context_data(**kwargs)
        context['tittle'] = "Seguimiento de Trámites"
        context['motto'] = "Bandeja de seguimiento - Historial"
        proceso_alumno = ProcesoAlumno.objects.filter(pk=self.kwargs['id']).first()
        context['process_alumno'] = proceso_alumno
        context['process_alumno_items'] = ProcesoAlumnoItems.objects.filter(process_alumno=proceso_alumno)
        context['historials'] = Historial.objects.filter(process_alumno_id=self.kwargs['id'])
        return context


## Stram files


class HistorialStreamFilesView(View):
    model = Historial

    def get_object(self):
        id = self.kwargs['id']
        if not self.model.objects.filter(id=id).exists():
            raise Http404
        historial = self.model.objects.get(id=id)
        return historial

    def get(self, request, *args, **kwargs):
        historial = self.get_object()
        if historial.document:
            document = historial.document.read()
            _split_ = os.path.splitext(str(historial.document))
            print "application/"+_split_[1][1:]
            return HttpResponse(document, content_type="application/"+_split_[1][1:])
        raise Http404


class ProcesoAlumnoStreamFilesView(View):
    model = ProcesoAlumno

    def get_object(self):
        id = self.kwargs['id']
        if not self.model.objects.filter(id=id).exists():
            raise Http404
        process = self.model.objects.get(id=id)
        return process

    def get(self, request, *args, **kwargs):
        process = self.get_object()
        if process.document:
            document = process.document.read()
            _split_ = os.path.splitext(str(process.document))
            print "application/"+_split_[1][1:]
            return HttpResponse(document, content_type="application/"+_split_[1][1:])
        raise Http404


class ProcesoAlumnoItemStreamFilesView(View):
    model = ProcesoAlumnoItems

    def get_object(self):
        id = self.kwargs['id']
        if not self.model.objects.filter(id=id).exists():
            raise Http404
        item = self.model.objects.get(id=id)
        return item

    def get(self, request, *args, **kwargs):
        item = self.get_object()
        if item.document:
            document = item.document.read()
            _split_ = os.path.splitext(str(item.document))
            print "application/"+_split_[1][1:]
            return HttpResponse(document, content_type="application/"+_split_[1][1:])
        raise Http404


# Alumno Crear Cuenta y Activación


class AlumnoNewCreateView(CreateView):
    form_class = AlumnoNewForm
    template_name = 'alumno/alumno_create.html'

    def send_mail_to_user(self, alumno, request):
        link = reverse_lazy('alumno-activate', kwargs={'token':alumno.token, 'username':alumno.user.username})
        context = {
            'email': alumno.user.email,
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'alumno': alumno,
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'],link),
            'protocol': 'https',
        }
        subject_template_name = 'alumno/new_account_subject.txt'
        email_template_name = 'alumno/new_account_activate.html'
        subject = loader.render_to_string(subject_template_name, context)
        subject = ''.join(subject.splitlines())
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=[alumno.user.email],
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)
    
    def get_success_url(self):
        return reverse('alumno-resend')

    def get(self, request, next=None):
        return super(AlumnoNewCreateView, self).get(request)

    def form_valid(self, form):
        try:
            if User.objects.filter(username=form.instance.user.username).exists():
                errors = form._errors.setdefault("email", ErrorList())
                errors.append(u"Ya existe un usuario registrado con este email")
                response = super(AlumnoNewCreateView, self).form_invalid(form)
            else:
                response = super(AlumnoNewCreateView, self).form_valid(form)
                _user_id_ = form.instance.user.id
                user = User.objects.filter(pk=_user_id_).first()
                alumno = Alumno.objects.filter(user=user).first()
                token = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(12))
                alumno.token = token
                alumno.save()
                self.send_mail_to_user(alumno, self.request)
        except Exception, e:
            print 'Error', e
            errors = form._errors.setdefault("email", ErrorList())
            errors.append(u"Ya existe un usuario registrado con este email")
            if form.instance.user.id:
                form.instance.user.delete()
            response = super(AlumnoNewCreateView, self).form_invalid(form)
        return response

    def form_invalid(self, form):
        return super(AlumnoNewCreateView, self).form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super(AlumnoNewCreateView, self).get_context_data(**kwargs)
        context['tittle'] = "Alumnos"
        context['motto'] = "Formulario para crear un nuevo Alumno"
        return context


class AlumnoActivateView(ListView):
    model = Alumno
    template_name = 'alumno/alumno_activate.html'

    def send_mail_to_user(self, alumno, request):
        link = reverse_lazy('alumno-activate', kwargs={'token':alumno.token, 'username':alumno.user.username})
        context = {
            'email': alumno.user.email,
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'alumno': alumno,
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'],link),
            'protocol': 'https',
        }
        subject_template_name = 'alumno/new_account_subject.txt'
        email_template_name = 'alumno/new_account_activate.html'
        subject = loader.render_to_string(subject_template_name, context)
        subject = ''.join(subject.splitlines())
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=[alumno.user.email],
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def get(self, request, *args, **kwargs):
        username = kwargs.get('username', None)
        user = User.objects.filter(username=username).first()
        token = kwargs.get('token', None)
        try:
            alumno = Alumno.objects.get(user__username=username)
            if alumno.token == token:
                alumno.is_active = True
                alumno.save()
                login(request, user)
                next = request.GET.get('next')
                if next is not None:
                    return redirect(next)
                return redirect('/')
            else:
                token = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(12))
                alumno.token = token
                alumno.save()
                self.send_mail_to_user(alumno, self.request)
                raise Http404
        except Exception as e:
            raise Http404


class AlumnoResendView(TemplateView):
    alert_message = {'class': 'alert alert-warning', 'mod': 'Aviso:', 'message': u'Para poder ingresar al sistema debe de activar su cuenta, se ha enviado el link de activación al correo electrónico ingresado anteriormente. Porfavor revisar.'}
    template_name = 'alumno/alumno_resend.html'

    def send_mail_to_user(self, alumno, request):
        link = reverse_lazy('alumno-activate', kwargs={'token':alumno.token, 'username':alumno.user.username})
        context = {
            'email': alumno.user.email,
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'alumno': alumno,
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'],link),
            'protocol': 'https',
        }
        subject_template_name = 'alumno/new_account_subject.txt'
        email_template_name = 'alumno/new_account_activate.html'
        subject = loader.render_to_string(subject_template_name, context)
        subject = ''.join(subject.splitlines())
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=[alumno.user.email],
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def post(self, request):
        if request.method == "POST":
            email = request.POST.get("id_email", "")
            if email:
                user = User.objects.filter(email=email).first()
                if user:
                    alumno = Alumno.objects.filter(user=user).first()
                    token = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(12))
                    alumno.token = token
                    alumno.save()
                    self.send_mail_to_user(alumno, self.request)
                    self.alert_message = {'class': 'alert alert-success',
                    'mod': 'Envío exitoso:',
                    'message': u'El correo ha sido enviado.'}
                else:
                    self.alert_message = {'class': 'alert alert-danger',
                    'mod': 'Envio exitoso:',
                    'message': u'Ingrese un correo electrónico válido.'}
            else:
                self.alert_message = {'class': 'alert alert-danger',
                'mod': 'Error:',
                'message': u'Ingrese un correo electrónico.'}
            return self.get(request)

    def get_success_url(self):
        return "/"

    def get(self, request, next=None):
        return super(AlumnoResendView, self).get(request)

    def get_context_data(self, **kwargs):
        context = super(AlumnoResendView, self).get_context_data(**kwargs)
        context['alert_message'] = self.alert_message
        context['tittle'] = "Alumnos"
        context['motto'] = "Reenviar Correo Activación"
        return context


class AlumnoResetView(TemplateView):
    alert_message = None
    template_name = 'alumno/alumno_reset.html'

    def send_mail_to_user(self, user, token, request):
        link = reverse('administrador-login')
        context = {
            'email': user.email,
            'domain': "%s%s"%('http://', request.META['HTTP_HOST']),
            'site_name': 'Seguimiento UCSG',
            'user': user,
            'token': token,
            'link': "%s%s%s"%('http://', request.META['HTTP_HOST'],link),
            'protocol': 'https',
        }
        email_template_name = 'alumno/account_reset.html'
        subject = 'Recuperar Contraseña'
        email = loader.get_template(email_template_name).render(context)
        email_message = EmailMessage(
            subject,
            email,
            to=[user.email],
            from_email=settings.DEFAULT_FROM_EMAIL
        )
        email_message.content_subtype = 'html'
        email_message.send(fail_silently=False)

    def post(self, request):
        if request.method == "POST":
            email = request.POST.get("id_email", "")
            if email:
                user = User.objects.filter(email=email).first()
                if user:
                    token = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(24))
                    user.set_password(token)
                    user.save()
                    self.send_mail_to_user(user, token, self.request)
                    self.alert_message = {'class': 'alert alert-success',
                    'mod': 'Cambio exitoso:',
                    'message': u'Se ha enviado un correo con tu nueva contraseña.'}
                else:
                    self.alert_message = {'class': 'alert alert-danger',
                    'mod': 'Envio exitoso:',
                    'message': u'Ingrese un correo electrónico válido.'}
            else:
                self.alert_message = {'class': 'alert alert-danger',
                'mod': 'Error:',
                'message': u'Ingrese un correo electrónico.'}
            return self.get(request)

    def get_success_url(self):
        return "/"

    def get(self, request, next=None):
        return super(AlumnoResetView, self).get(request)

    def get_context_data(self, **kwargs):
        context = super(AlumnoResetView, self).get_context_data(**kwargs)
        context['alert_message'] = self.alert_message
        context['tittle'] = "Recuperar Contraseña"
        context['motto'] = "Recuperar Contraseña"
        return context

